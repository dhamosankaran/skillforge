"""Resume parser supporting PDF and DOCX formats."""
import io
import re
from typing import Any, Dict, List, Optional

import pdfplumber
from docx import Document

from app.utils.text_cleaner import (
    clean_resume_text,
    extract_email,
    extract_phone,
    extract_urls,
    normalize_whitespace,
)

# Common section header patterns
SECTION_PATTERNS = {
    "summary": r"(?i)^(summary|professional summary|objective|profile|about me|overview)",
    "experience": r"(?i)^(experience|work experience|employment|professional experience|work history|career history)",
    "education": r"(?i)^(education|academic background|qualifications|academic history|degrees)",
    "skills": r"(?i)^(skills|technical skills|core competencies|competencies|technologies|tech stack|expertise)",
    "projects": r"(?i)^(projects|personal projects|key projects|notable projects|portfolio)",
    "certifications": r"(?i)^(certifications|certificates|licenses|credentials)",
    "awards": r"(?i)^(awards|honors|achievements|accomplishments|recognition)",
    "publications": r"(?i)^(publications|papers|research|articles)",
    "languages": r"(?i)^(languages|spoken languages|foreign languages)",
    "volunteer": r"(?i)^(volunteer|volunteering|community|extracurricular)",
}

BULLET_PATTERNS = re.compile(
    r"^[\•\-\*\>\◦\▪\▸\►\◼\●\·\○\—\–]\s+(.+)", re.MULTILINE
)


def parse_pdf(file_bytes: bytes) -> Dict[str, Any]:
    """Parse a PDF resume and return structured data.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Dictionary with full_text, sections, bullet_points, contact_info.
    """
    text_parts: List[str] = []
    formatting_hints: Dict[str, bool] = {
        "has_tables": False,
        "has_images": False,
        "multi_column": False,
    }

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # Detect tables
            tables = page.find_tables()
            if tables:
                formatting_hints["has_tables"] = True

            # Detect images
            if page.images:
                formatting_hints["has_images"] = True

            # Detect multi-column layout by checking word bounding boxes
            words = page.extract_words()
            if words:
                x_positions = [w["x0"] for w in words]
                if x_positions:
                    # If there are words starting far to the right, likely multi-column
                    page_width = page.width
                    right_column_words = [x for x in x_positions if x > page_width * 0.5]
                    left_column_words = [x for x in x_positions if x < page_width * 0.4]
                    if len(right_column_words) > 10 and len(left_column_words) > 10:
                        formatting_hints["multi_column"] = True

            page_text = page.extract_text() or ""
            text_parts.append(page_text)

    full_text = "\n".join(text_parts)
    full_text = clean_resume_text(full_text)

    return {
        "full_text": full_text,
        "sections": detect_sections(full_text),
        "bullet_points": extract_bullets(full_text),
        "contact_info": extract_contact_info(full_text),
        "formatting_hints": formatting_hints,
        "source_type": "pdf",
    }


def parse_docx(file_bytes: bytes) -> Dict[str, Any]:
    """Parse a DOCX resume and return structured data.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Dictionary with full_text, sections, bullet_points, contact_info.
    """
    doc = Document(io.BytesIO(file_bytes))
    formatting_hints: Dict[str, bool] = {
        "has_tables": False,
        "has_images": False,
        "multi_column": False,
    }

    # Check for tables
    if doc.tables:
        formatting_hints["has_tables"] = True

    # Check for inline images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            formatting_hints["has_images"] = True
            break

    # Check for multi-column sections
    for section in doc.sections:
        if section.start_type is not None:
            cols = section._sectPr.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols"
            )
            if cols:
                num_cols = cols[0].get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num", "1"
                )
                if int(num_cols) > 1:
                    formatting_hints["multi_column"] = True

    # Extract all paragraph text
    paragraphs: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)

    full_text = "\n".join(paragraphs)
    full_text = clean_resume_text(full_text)

    return {
        "full_text": full_text,
        "sections": detect_sections(full_text),
        "bullet_points": extract_bullets(full_text),
        "contact_info": extract_contact_info(full_text),
        "formatting_hints": formatting_hints,
        "source_type": "docx",
    }


def detect_sections(text: str) -> Dict[str, str]:
    """Detect and extract resume sections from plain text.

    Args:
        text: Plain text of the resume.

    Returns:
        Dictionary mapping section names to their content.
    """
    lines = text.split("\n")
    sections: Dict[str, str] = {}
    current_section: Optional[str] = None
    current_content: List[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Check if this line is a section header
        detected_section = None
        for section_name, pattern in SECTION_PATTERNS.items():
            if re.match(pattern, stripped):
                detected_section = section_name
                break

        if detected_section:
            # Save previous section
            if current_section and current_content:
                sections[current_section] = "\n".join(current_content).strip()
            current_section = detected_section
            current_content = []
        else:
            if current_section:
                current_content.append(stripped)

    # Save the last section
    if current_section and current_content:
        sections[current_section] = "\n".join(current_content).strip()

    return sections


def extract_bullets(text: str) -> List[str]:
    """Extract bullet points from resume text.

    Args:
        text: Plain text of the resume.

    Returns:
        List of bullet point strings (without bullet characters).
    """
    bullets: List[str] = []

    # Match explicit bullet characters
    for match in BULLET_PATTERNS.finditer(text):
        bullet_text = match.group(1).strip()
        if len(bullet_text) > 10:  # Filter out very short non-bullets
            bullets.append(bullet_text)

    # If no explicit bullets found, try to extract achievement-style sentences
    if not bullets:
        lines = text.split("\n")
        for line in lines:
            stripped = line.strip()
            # Lines that start with action verbs and are substantial
            if len(stripped) > 30 and re.match(r"^[A-Z][a-z]+ed|^[A-Z][a-z]+ed|^[A-Z][a-z]+ing", stripped):
                bullets.append(stripped)

    return bullets[:50]  # Cap at 50 bullets


def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract contact information from resume text.

    Args:
        text: Plain text of the resume.

    Returns:
        Dictionary with email, phone, and urls.
    """
    # Try to get name from first non-empty line
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    potential_name = lines[0] if lines else ""
    # If first line looks like a name (2-4 words, no special chars)
    if potential_name and len(potential_name.split()) in (2, 3, 4):
        if re.match(r"^[A-Za-z\s\.\-]+$", potential_name):
            name = potential_name
        else:
            name = ""
    else:
        name = ""

    return {
        "name": name,
        "email": extract_email(text),
        "phone": extract_phone(text),
        "urls": ", ".join(extract_urls(text)),
    }
